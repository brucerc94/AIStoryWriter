"""
Export the compiled novel to DOCX or PDF.

Uses python-docx and reportlab — both optional, local-only libraries (no
network, no cloud). If not installed, the export functions raise a clear
RuntimeError with an install hint that the UI surfaces to the user.
"""

from __future__ import annotations

from engine.models import Project


def _paragraphs(text: str) -> list[str]:
    """Split raw chapter/synopsis text into paragraphs on blank lines."""
    if not text:
        return []
    return [p.strip() for p in text.split("\n\n") if p.strip()]


def export_to_docx(project: Project, output_path: str) -> None:
    """Write the full novel (title, synopsis, chapters) to a .docx file."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from e

    doc = Document()

    doc.add_heading(project.title, level=0)

    if project.synopsis.strip():
        doc.add_heading("Synopsis", level=1)
        for para in _paragraphs(project.synopsis):
            doc.add_paragraph(para)
        doc.add_page_break()

    for ch in sorted(project.chapters, key=lambda c: c.number):
        doc.add_heading(f"Chapter {ch.number}: {ch.title}", level=1)
        paras = _paragraphs(ch.content)
        if not paras:
            doc.add_paragraph("(This chapter hasn't been written yet.)")
        for para in paras:
            doc.add_paragraph(para)
        doc.add_page_break()

    doc.save(output_path)


def export_to_pdf(project: Project, output_path: str) -> None:
    """Write the full novel (title, synopsis, chapters) to a .pdf file."""
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            PageBreak,
        )
    except ImportError as e:
        raise RuntimeError(
            "reportlab is not installed. Run: pip install reportlab"
        ) from e

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    heading_style = styles["Heading1"]
    body_style = ParagraphStyle(
        "StoryBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        spaceAfter=10,
    )

    doc = SimpleDocTemplate(
        output_path,
        pagesize=LETTER,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        title=project.title,
    )

    flow = [Paragraph(_escape(project.title), title_style), Spacer(1, 24)]

    if project.synopsis.strip():
        flow.append(Paragraph("Synopsis", heading_style))
        for para in _paragraphs(project.synopsis):
            flow.append(Paragraph(_escape(para), body_style))
        flow.append(PageBreak())

    for ch in sorted(project.chapters, key=lambda c: c.number):
        flow.append(Paragraph(f"Chapter {ch.number}: {_escape(ch.title)}", heading_style))
        paras = _paragraphs(ch.content)
        if not paras:
            flow.append(Paragraph("(This chapter hasn't been written yet.)", body_style))
        for para in paras:
            flow.append(Paragraph(_escape(para), body_style))
        flow.append(PageBreak())

    doc.build(flow)


def _escape(text: str) -> str:
    """Escape for reportlab's mini-HTML Paragraph markup, keep line breaks."""
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    return escaped.replace("\n", "<br/>")
